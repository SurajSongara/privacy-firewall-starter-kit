"""Convert non-PDF source files (images, txt, md, docx) to PDF.

The detection pipeline is PDF-native, so other formats are converted to
PDF once at ingestion and the converted file is fed through the existing
pipeline unchanged. Conversions are pure PyMuPDF except DOCX, which
needs the optional ``python-docx`` package for text extraction.
"""

from __future__ import annotations

from pathlib import Path

import fitz

IMAGE_SUFFIXES: frozenset[str] = frozenset(
    {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".gif"}
)
"""Raster image formats PyMuPDF can open and convert directly."""

TEXT_SUFFIXES: frozenset[str] = frozenset({".txt", ".md"})
"""Plain-text formats rendered onto PDF pages as-is."""

DOCX_SUFFIXES: frozenset[str] = frozenset({".docx"})
"""Word documents (text extracted via the optional ``python-docx``)."""

SUPPORTED_SUFFIXES: frozenset[str] = frozenset(
    {".pdf"} | IMAGE_SUFFIXES | TEXT_SUFFIXES | DOCX_SUFFIXES
)
"""Every file suffix the studio/ingestion layer accepts."""

_PAGE_WIDTH = 595.0  # A4 in points
_PAGE_HEIGHT = 842.0
_MARGIN = 50.0
_FONT_NAME = "cour"  # monospace: predictable wrapping and value alignment
_FONT_SIZE = 10.0
_LINE_HEIGHT = _FONT_SIZE * 1.4


class ConversionError(ValueError):
    """A source file could not be converted to PDF."""


def is_supported(path: Path | str) -> bool:
    """Whether *path*'s suffix is an accepted document format."""
    return Path(path).suffix.lower() in SUPPORTED_SUFFIXES


def needs_conversion(path: Path | str) -> bool:
    """Whether *path* is a supported format that must be converted first."""
    suffix = Path(path).suffix.lower()
    return suffix in SUPPORTED_SUFFIXES and suffix != ".pdf"


def convert_to_pdf(source: Path, dest: Path) -> Path:
    """Convert *source* to a PDF at *dest* (cached by modification time).

    Args:
        source: The input file (image, txt, md, or docx).
        dest: Where to write the converted PDF.

    Returns:
        *dest*, for chaining.

    Raises:
        ConversionError: If the format is unsupported, the file is
            unreadable/corrupt, or DOCX support is not installed.
    """
    source = Path(source)
    dest = Path(dest)
    suffix = source.suffix.lower()
    if not source.exists():
        msg = f"source file not found: {source}"
        raise ConversionError(msg)
    if dest.exists() and dest.stat().st_mtime >= source.stat().st_mtime:
        return dest  # up-to-date conversion already on disk

    if suffix in IMAGE_SUFFIXES:
        _image_to_pdf(source, dest)
    elif suffix in TEXT_SUFFIXES:
        text = source.read_text(encoding="utf-8", errors="replace")
        _text_to_pdf(text, dest)
    elif suffix in DOCX_SUFFIXES:
        _text_to_pdf(_extract_docx_text(source), dest)
    else:
        msg = f"unsupported file type: {suffix or '(no extension)'}"
        raise ConversionError(msg)
    return dest


def _image_to_pdf(source: Path, dest: Path) -> None:
    """Wrap a raster image into a single-page PDF (no text layer — OCR's job)."""
    try:
        with fitz.open(str(source)) as img:
            pdf_bytes = img.convert_to_pdf()
    except Exception as exc:
        msg = f"could not read image {source.name}: {exc}"
        raise ConversionError(msg) from exc
    dest.write_bytes(pdf_bytes)


def _text_to_pdf(text: str, dest: Path) -> None:
    """Render plain text onto paginated A4 pages with a real text layer."""
    max_width = _PAGE_WIDTH - 2 * _MARGIN
    char_width = fitz.get_text_length("M", fontname=_FONT_NAME, fontsize=_FONT_SIZE)
    chars_per_line = max(1, int(max_width / char_width))
    lines_per_page = max(1, int((_PAGE_HEIGHT - 2 * _MARGIN) / _LINE_HEIGHT))

    lines: list[str] = []
    for raw in text.splitlines() or [""]:
        raw = raw.replace("\t", "    ")
        lines.extend(_wrap_line(raw, chars_per_line))

    doc = fitz.open()
    try:
        for start in range(0, max(len(lines), 1), lines_per_page):
            page = doc.new_page(width=_PAGE_WIDTH, height=_PAGE_HEIGHT)
            y = _MARGIN + _FONT_SIZE
            for line in lines[start : start + lines_per_page]:
                if line:
                    page.insert_text((_MARGIN, y), line, fontsize=_FONT_SIZE, fontname=_FONT_NAME)
                y += _LINE_HEIGHT
        doc.save(str(dest))
    finally:
        doc.close()


def _wrap_line(line: str, width: int) -> list[str]:
    """Wrap one logical line at *width* characters, breaking on spaces."""
    if len(line) <= width:
        return [line]
    wrapped: list[str] = []
    while len(line) > width:
        cut = line.rfind(" ", 1, width + 1)
        if cut <= 0:
            cut = width
        wrapped.append(line[:cut])
        line = line[cut:].lstrip(" ")
    wrapped.append(line)
    return wrapped


def _extract_docx_text(source: Path) -> str:
    """Pull paragraph and table text out of a DOCX file.

    Raises:
        ConversionError: If ``python-docx`` is missing or the file is
            not a valid DOCX document.
    """
    try:
        import docx
    except ImportError as exc:
        msg = "DOCX support requires the python-docx package: pip install python-docx"
        raise ConversionError(msg) from exc

    try:
        document = docx.Document(str(source))
    except Exception as exc:
        msg = f"could not read DOCX {source.name}: {exc}"
        raise ConversionError(msg) from exc

    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
